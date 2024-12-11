# backend/app/loader.py

import os
import tempfile
import logging
from uuid import uuid4
from urllib.parse import urlparse
from app.config import settings
from app.db import get_weaviate_client
from langchain.docstore.document import Document
from app.utils.splitter import TextSplitter
from langchain.document_loaders import PyPDFLoader, UnstructuredWordDocumentLoader
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain_weaviate.vectorstores import WeaviateVectorStore
from pdfminer.high_level import extract_text
import docx
from urllib.parse import urlparse, unquote
from fastapi import FastAPI
from app.models import DocumentOut  # If needed
logger = logging.getLogger(__name__)

async def download_file(url: str) -> bytes:
    """
    Downloads a file from a given URL.
    """
    import aiohttp
    async with aiohttp.ClientSession() as session:
        async with session.get(url) as resp:
            if resp.status == 200:
                return await resp.read()
            else:
                raise RuntimeError(f"Failed to download file from {url} (Status: {resp.status})")

async def ingest_and_index(file_url: str, upload_id: str,app: FastAPI):
    try:
        # Download the file
        logger.info(f"Downloading file from {file_url}...")
        file_content = await download_file(file_url)

        # Parse the URL to get the filename
        parsed_url = urlparse(file_url)
        path = parsed_url.path
        filename = os.path.basename(path)
        filename = unquote(filename)
        suffix = os.path.splitext(filename)[1].lower()
        logger.debug(f"Determined file suffix: {suffix}")

        if suffix not in ['.pdf', '.docx']:
            raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")

        # Save to a temporary file for processing
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name

        # Load the document
        logger.info(f"Loading document from {temp_file_path}...")
        if suffix == ".pdf":
            text = extract_text(temp_file_path)
        elif suffix == ".docx":
            from io import BytesIO
            doc = docx.Document(temp_file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")

        # Remove the temporary file
        os.remove(temp_file_path)

        # Split the document into chunks
        logger.info("Splitting the document into chunks using custom TextSplitter...")
        text_splitter = TextSplitter(chunk_size=512, chunk_overlap=20)
        chunks = text_splitter.split(text)
        logger.info(f"Split the document into {len(chunks)} chunks.")

        # Convert chunks into LangChain Document objects with `upload_id` in metadata
        documents = [Document(page_content=chunk, metadata={"source": file_url, "upload_id": upload_id}) for chunk in chunks]

        # Generate embeddings for the chunks
        logger.info("Generating embeddings for document chunks...")
        embeddings = OpenAIEmbeddings(openai_api_key=settings.OPENAI_API_KEY)

        # Connect to Weaviate
        logger.info("Connecting to Weaviate...")
        client = get_weaviate_client(app)

        # Index the documents into Weaviate
        logger.info("Indexing document chunks into Weaviate...")
        vectorstore = WeaviateVectorStore.from_documents(
            documents,
            embeddings,
            client=client,
            index_name="ChatDocument",
            text_key="content",
        )

        logger.info("Document successfully ingested and indexed into Weaviate.")

    except Exception as e:
        logger.error(f"Failed to ingest and index the document: {e}")
        raise RuntimeError(f"Failed to ingest and index the document: {e}")


